#!/usr/bin/env python3
"""
Resume Formatter - Converts markdown resume to properly styled DOCX
Applies the standard Isaac Kehle resume formatting:
- Blue (#2E5FA3) headers, Gray (#444444) dates/contact
- Font sizes: 18pt name, 11pt subtitle, 13pt section headers, 10.5pt company headers, 10pt job titles, 9.5pt bullets, 9pt contact
- Company format: COMPANY (Contract) | Location | Dates
- Margins: 0.75" top/bottom, 1" left/right
"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colors
BLUE = RGBColor(0x2E, 0x5F, 0xA3)
GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_paragraph_style(
    para,
    font_size=Pt(9.5),
    bold=False,
    italic=False,
    color=BLACK,
    align=None,
    underline=False,
):
    """Apply consistent paragraph styling."""
    run = para.runs[0] if para.runs else para.add_run()
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Arial"
    run.font.underline = underline
    if align:
        para.alignment = align
    return para


def add_inline_runs(para, text, font_size=Pt(9.5), color=BLACK, base_bold=False):
    """Add runs to para, honouring **bold** markers anywhere in text."""
    # Split on ** — odd-indexed segments are bold spans
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = font_size
        run.font.color.rgb = color
        run.font.name = "Arial"
        # Even index = normal, odd index = bold (or inverted if base_bold)
        run.font.bold = base_bold if i % 2 == 0 else (not base_bold)


def format_document(doc):
    """Apply document-wide formatting."""
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_height = Inches(11.0)
        section.page_width = Inches(8.5)


def parse_markdown(md_content):
    """Parse markdown and create structured document."""
    lines = md_content.split("\n")
    sections = []
    current_section = None
    current_items = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines at start
        if not stripped and not current_section:
            continue

        # Name header (# NAME)
        if stripped.startswith("# ") and not current_section:
            current_section = "NAME"
            current_items = [stripped]
            continue

        # Section headers (## HEADER)
        if stripped.startswith("## "):
            if current_section:
                sections.append((current_section, current_items))
            current_section = stripped[3:].strip()
            current_items = []
            continue

        # Subsection headers (### HEADER) - but NOT job headers with pipes " | "
        if stripped.startswith("### "):
            job_candidate = stripped[4:].strip()
            # Job headers have " | " (company name | location | dates), treat as items not sections
            if " | " in job_candidate:
                current_items.append(stripped)
                continue
            else:
                # Regular subsection header
                if current_section:
                    sections.append((current_section, current_items))
                current_section = job_candidate
                current_items = []
                continue

        # Horizontal rules (skip, don't use as section delimiter)
        if stripped.startswith("---"):
            continue

        # Table rows (start with |)
        if stripped.startswith("|"):
# In parse_markdown, don't skip separator rows but treat them as regular rows
                # Skip separator rows (e.g., |---|---|)
                # if all(c in "|-: " for c in stripped):
                #     continue
                # Instead, process all | lines including separators

            current_items.append("TABLE|" + stripped)
            continue

        # Regular content
        if stripped:
            current_items.append(stripped)

    if current_section:
        sections.append((current_section, current_items))

    return sections


def build_resume(sections, doc):
    """Build the resume document from parsed sections."""

    # Process name section (handled separately as NAME)
    name_section = None
    for section_name, items in sections:
        if section_name == "NAME":
            name_section = items
            break

    if name_section:
        # Name (# ISAAC KEHLE)
        name = name_section[0].replace("#", "").strip()
        para = doc.add_paragraph(name)
        set_paragraph_style(
            para,
            font_size=Pt(18),
            bold=True,
            color=BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        # Subtitle (**title**) - blue and bold like master
        if len(name_section) > 1:
            subtitle = name_section[1].replace("**", "").strip()
            para = doc.add_paragraph(subtitle)
            set_paragraph_style(
                para,
                font_size=Pt(11),
                bold=True,
                color=BLUE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

        # Contact info (line with | separators) - preserve angle brackets around email
        for item in name_section[2:]:
            if "|" in item or "@" in item or "http" in item:
                # Don't strip angle brackets - keep <email> format
                para = doc.add_paragraph(item.replace("**", "").strip())
                set_paragraph_style(
                    para, font_size=Pt(9), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER
                )

    # Process remaining sections
    for section_name, items in sections:
        if section_name == "---" or not section_name:
            continue

        if section_name == "NAME":
            continue  # Already processed

        if section_name.startswith("ISAAC"):
            continue  # Already processed

        # Section header with underline
        para = doc.add_paragraph(section_name.upper())
        # Use smaller size (10pt) for job-like headers like "ADDITIONAL EARLY-CAREER ROLES"
        if (
            "ADDITIONAL" in section_name.upper()
            or "EARLY-CAREER" in section_name.upper()
        ):
            set_paragraph_style(para, font_size=Pt(10), bold=True, color=BLACK)
        else:
            set_paragraph_style(
                para, font_size=Pt(13), bold=True, color=BLUE, underline=True
            )
        para.paragraph_format.space_before = Pt(6)

        # Section content
        # First, group consecutive TABLE| items into table blocks
        grouped = []
        i = 0
        while i < len(items):
            if items[i] is not None and items[i].startswith("TABLE|"):
                table_rows = []
                while i < len(items) and items[i] is not None and items[i].startswith("TABLE|"):
                    raw = items[i][6:]  # Strip "TABLE|" prefix
                    cells = [c.strip() for c in raw.split("|")][1:-1]
                    table_rows.append(cells)
                    i += 1
                grouped.append(("__TABLE__", table_rows))
            else:
                grouped.append(("__TEXT__", items[i] if items[i] is not None else ""))
                i += 1

        for item_type, item_value in grouped:
            if item_type == "__TABLE__":
                table_rows = item_value
                if not table_rows:
                    continue
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                # Fixed layout for consistent column widths
                table.autofit = False

                # Set table width to 100% of page (modify existing tblW if present)
                tbl_pr = table._tbl.tblPr
                existing_tbl_w = tbl_pr.find(qn("w:tblW"))
                if existing_tbl_w is not None:
                    existing_tbl_w.set(qn("w:w"), "5000")
                    existing_tbl_w.set(qn("w:type"), "pct")
                else:
                    tbl_w = OxmlElement("w:tblW")
                    tbl_w.set(qn("w:w"), "5000")
                    tbl_w.set(qn("w:type"), "pct")
                    tbl_pr.append(tbl_w)

                # Remove all borders via XML
                borders = OxmlElement("w:tblBorders")
                for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = OxmlElement(f"w:{border_name}")
                    el.set(qn("w:val"), "nil")
                    borders.append(el)
                tbl_pr.append(borders)

                # Set column widths based on column count
                if num_cols == 2:
                    col_widths = [Inches(2.0), Inches(4.5)]
                elif num_cols == 3:
                    col_widths = [Inches(2.5), Inches(2.5), Inches(1.5)]
                elif num_cols == 4:
                    col_widths = [Inches(2.0), Inches(2.0), Inches(1.3), Inches(1.2)]
                else:
                    col_widths = [Inches(6.5 / num_cols)] * num_cols

                # Set grid column widths to match (required for fixed layout)
                for j, width in enumerate(col_widths):
                    dxa = int(width / 635)
                    grid_col = table.columns[j]._gridCol
                    grid_col.set(qn("w:w"), str(dxa))
                    grid_col.set(qn("w:type"), "dxa")

                for i, row_cells in enumerate(table_rows):
                    for j, cell_text in enumerate(row_cells):
                        if j >= num_cols:
                            break
                        cell = table.cell(i, j)
                        # Set column width
                        if j < len(col_widths):
                            cell.width = col_widths[j]
                        # Clear default paragraph and write
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(cell_text)
                        run.font.size = Pt(9)
                        run.font.name = "Arial"
                        run.font.color.rgb = BLACK
                        if i == 0:
                            run.font.bold = True
                        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

                # Add thin bottom border to header row
                for cell in table.rows[0].cells:
                    tc_pr = cell._tc.tcPr
                    tc_borders = OxmlElement("w:tcBorders")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "8")
                    bottom.set(qn("w:color"), "666666")
                    tc_borders.append(bottom)
                    tc_pr.append(tc_borders)

                continue

            item = item_value
            # Job header (### style in content)
            if item.startswith("### "):
                job_line = item[4:].strip()
                if " | " in job_line:
                    # Parse: COMPANY (Contract) | Location | Dates
                    parts = job_line.split(" | ")
                    para = doc.add_paragraph()

                    # Company name (may include contract info in parentheses)
                    # Font: 10.5pt, bold for company name, regular for contract info
                    company_part = parts[0].strip()
                    if "(" in company_part and ")" in company_part:
                        # Split company name and contract info
                        paren_idx = company_part.find("(")
                        company_name = company_part[:paren_idx].strip()
                        contract_info = company_part[paren_idx:].strip()

                        # Company name - bold, 10.5pt
                        company_run = para.add_run(company_name)
                        company_run.font.size = Pt(10.5)
                        company_run.font.bold = True
                        company_run.font.color.rgb = BLACK
                        company_run.font.name = "Arial"

                        # Space before contract info
                        space_run = para.add_run(" ")
                        space_run.font.size = Pt(10.5)
                        space_run.font.name = "Arial"

                        # Contract info - regular, 10.5pt
                        contract_run = para.add_run(contract_info)
                        contract_run.font.size = Pt(10.5)
                        contract_run.font.bold = False
                        contract_run.font.color.rgb = BLACK
                        contract_run.font.name = "Arial"
                    else:
                        # Just company name, no contract info
                        company_run = para.add_run(company_part)
                        company_run.font.size = Pt(10.5)
                        company_run.font.bold = True
                        company_run.font.color.rgb = BLACK
                        company_run.font.name = "Arial"

                    # Pipe separator
                    pipe_run = para.add_run(" | ")
                    pipe_run.font.size = Pt(10.5)
                    pipe_run.font.color.rgb = BLACK
                    pipe_run.font.name = "Arial"

                    # Location and dates - regular, 10.5pt
                    detail_text = (
                        " | ".join(parts[1:]).strip() if len(parts) > 1 else ""
                    )
                    if detail_text:
                        detail_run = para.add_run(detail_text)
                        detail_run.font.size = Pt(10.5)
                        detail_run.font.bold = False
                        detail_run.font.color.rgb = BLACK
                        detail_run.font.name = "Arial"

                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(2)
                else:
                    para = doc.add_paragraph(job_line)
                    set_paragraph_style(para, font_size=Pt(10), bold=True, color=BLACK)
                    para.paragraph_format.space_after = Pt(4)

            # Bold role title under job
            elif item.startswith("**") and item.endswith("**"):
                para = doc.add_paragraph(item.replace("**", ""))
                set_paragraph_style(para, font_size=Pt(10), bold=True, color=BLACK)
                para.paragraph_format.space_after = Pt(4)

            # Bullet points (with inline **bold** support anywhere)
            elif item.startswith("- "):
                bullet_text = item[2:]
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(4)
                add_inline_runs(para, bullet_text, font_size=Pt(9.5), color=BLACK)

            # Regular bold lines (Education, Clearance, etc.) — whole line starts **bold**
            elif item.startswith("**") and not item.startswith("- "):
                para = doc.add_paragraph()
                add_inline_runs(para, item, font_size=Pt(9.5), color=BLACK)
                para.paragraph_format.space_after = Pt(4)

            # Italic lines
            elif item.startswith("*") and item.endswith("*"):
                para = doc.add_paragraph(item.replace("*", ""))
                set_paragraph_style(para, font_size=Pt(9), italic=True, color=GRAY)
                para.paragraph_format.space_after = Pt(4)

            # URLs
            elif item.startswith("http"):
                para = doc.add_paragraph(item)
                set_paragraph_style(para, font_size=Pt(9), color=GRAY)
                para.paragraph_format.space_after = Pt(4)

            # Technical expertise lines (Key: Value)
            elif ":" in item and item.split(":")[0].count("**") == 2:
                parts = item.split(":", 1)
                if len(parts) == 2:
                    para = doc.add_paragraph()
                    key_run = para.add_run(parts[0].replace("**", "") + ":")
                    key_run.font.size = Pt(9.5)
                    key_run.font.bold = True
                    key_run.font.color.rgb = BLACK
                    key_run.font.name = "Arial"
                    val_run = para.add_run(parts[1].strip())
                    val_run.font.size = Pt(9.5)
                    val_run.font.color.rgb = BLACK
                    val_run.font.name = "Arial"
                    para.paragraph_format.space_after = Pt(4)

            # Default text (supports inline **bold**)
            else:
                para = doc.add_paragraph()
                add_inline_runs(para, item, font_size=Pt(9.5), color=BLACK)
                para.paragraph_format.space_after = Pt(4)


def main():
    if len(sys.argv) != 3:
        print("Usage: resume-formatter <input.md> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Read markdown
    with open(input_file, "r") as f:
        md_content = f.read()

    # Parse and build
    sections = parse_markdown(md_content)
    doc = Document()
    format_document(doc)
    build_resume(sections, doc)

    # Save
    doc.save(output_file)
    print(f"Generated: {output_file}")


if __name__ == "__main__":
    main()
