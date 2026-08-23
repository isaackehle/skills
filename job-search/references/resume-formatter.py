#!/usr/bin/env python3
"""
Resume Markdown Exporter - Converts markdown resume to properly styled DOCX
Applies the Calibri-based Isaac Kehle resume formatting:
- Calibri font throughout
- Section header color: #1F3864 (dark blue) with bottom border
- Name: 17pt bold black (#000000), centered
- Role line: 10pt bold #1F3864, centered
- Contact bar: 9pt, #444444 with · separators, #1F3864 bottom border (sz 8)
- Section headers: 10.5pt bold #1F3864, bottom border (sz 6), 9pt before / 4.5pt after
- Job lines: company 10.5pt bold black, " — " 10pt, location 10pt italic, dates 9.5pt #444444
- Role titles: 10pt bold italic #1F3864
- Bullets: 10pt, List Bullet style, 3.5pt after
- Body text: 10pt black (#000000)
- Margins: 0.625" L/R, 0.43" T/B, US Letter
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x38, 0x64)  # #1F3864
GRAY = RGBColor(0x44, 0x44, 0x44)

# Font
FONT = 'Calibri'

def set_font(run, size=Pt(10), bold=False, italic=False, color=BLACK):
    """Apply font styling to a run."""
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def rgb_to_hex(rgb):
    """Convert RGBColor to hex string for docx."""
    # RGBColor stores values as _r, _g, _b internally
    r = getattr(rgb, '_r', rgb[0]) if hasattr(rgb, '_r') else rgb[0]
    g = getattr(rgb, '_g', rgb[1]) if hasattr(rgb, '_g') else rgb[1]
    b = getattr(rgb, '_b', rgb[2]) if hasattr(rgb, '_b') else rgb[2]
    return f'{r:02X}{g:02X}{b:02X}'

def add_bottom_border(para, color, size=Pt(6)):
    """Add bottom border to paragraph."""
    el = para._element
    pPr = el.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pb = OxmlElement('w:bottom')
    pb.set(qn('w:val'), 'single')
    # Convert points to half-points (docx uses half-points internally)
    # Pt(6) = 6 points = 12 half-points
    # Pt is a float subclass, so we can just multiply by 2
    half_points = int(float(size) * 2)
    pb.set(qn('w:sz'), str(half_points))
    pb.set(qn('w:color'), rgb_to_hex(color))
    pb.set(qn('w:space'), '0')
    pBdr.append(pb)
    pPr.append(pBdr)

def parse_markdown(md_content):
    """Parse markdown and create structured document."""
    lines = md_content.split('\n')
    sections = []
    current_section = None
    current_items = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines at start
        if not stripped and not current_section:
            continue

        # Name header (# NAME)
        if stripped.startswith('# ') and not current_section:
            current_section = 'NAME'
            current_items = [stripped]
            continue

        # Section headers (## HEADER)
        if stripped.startswith('## '):
            if current_section:
                sections.append((current_section, current_items))
            current_section = stripped[3:].strip()
            current_items = []
            continue

        # Subsection headers (### HEADER) - but NOT job headers with " | "
        if stripped.startswith('### '):
            job_candidate = stripped[4:].strip()
            # Job headers have " | " (company name | location | dates), treat as items not sections
            if ' | ' in job_candidate:
                current_items.append(stripped)
                continue
            else:
                # Regular subsection header
                if current_section:
                    sections.append((current_section, current_items))
                current_section = job_candidate
                current_items = []
                continue

        # Horizontal rules (skip, don't use as section delimiter)
        if stripped.startswith('---'):
            continue

        # Regular content
        if stripped:
            current_items.append(stripped)

    if current_section:
        sections.append((current_section, current_items))

    return sections

def build_resume(sections, doc):
    """Build the resume document from parsed sections."""

    # Process name section (handled separately as NAME)
    name_section = None
    for section_name, items in sections:
        if section_name == 'NAME':
            name_section = items
            break

    if name_section:
        # Name (# ISAAC KEHLE) - 17pt bold black, centered
        name = name_section[0].replace('#', '').strip()
        para = doc.add_paragraph(name)
        run = para.runs[0] if para.runs else para.add_run()
        set_font(run, size=Pt(17), bold=True, color=BLACK)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle (**title**) - 10pt bold #1F3864, centered
        if len(name_section) > 1:
            subtitle = name_section[1].replace('**', '').strip()
            para = doc.add_paragraph(subtitle)
            run = para.runs[0] if para.runs else para.add_run()
            set_font(run, size=Pt(10), bold=True, color=BLUE)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info (line with | separators) - 9pt, #444444 with · separators
        # Add bottom border in #1F3864 (sz 8)
        for item in name_section[2:]:
            if '|' in item or '@' in item or 'http' in item:
                # Replace | with ·
                contact_text = item.replace('|', ' · ').replace('**', '').strip()
                para = doc.add_paragraph(contact_text)
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(9), color=GRAY)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_bottom_border(para, BLUE, Pt(8))

    # Process remaining sections
    for section_name, items in sections:
        if section_name == '---' or not section_name:
            continue

        if section_name == 'NAME':
            continue  # Already processed

        if section_name.startswith('ISAAC'):
            continue  # Already processed

        # Section header with bottom border - 10.5pt bold #1F3864
        para = doc.add_paragraph(section_name.upper())
        run = para.runs[0] if para.runs else para.add_run()
        set_font(run, size=Pt(10.5), bold=True, color=BLUE)
        add_bottom_border(para, BLUE, Pt(6))
        para.paragraph_format.space_before = Pt(9)
        para.paragraph_format.space_after = Pt(4.5)

        # Section content
        for item in items:
            # Job header (### style in content)
            if item.startswith('### '):
                job_line = item[4:].strip()
                if ' | ' in job_line:
                    # Parse: COMPANY (Contract) | Location | Dates
                    parts = job_line.split(' | ')
                    para = doc.add_paragraph()

                    # Company name - 10.5pt bold black
                    company_part = parts[0].strip()
                    if '(' in company_part and ')' in company_part:
                        # Split company name and contract info
                        paren_idx = company_part.find('(')
                        company_name = company_part[:paren_idx].strip()
                        contract_info = company_part[paren_idx:].strip()

                        # Company name - bold, 10.5pt
                        company_run = para.add_run(company_name)
                        set_font(company_run, size=Pt(10.5), bold=True, color=BLACK)

                        # Space before contract info
                        space_run = para.add_run(' ')
                        set_font(space_run, size=Pt(10.5))

                        # Contract info - regular, 10.5pt
                        contract_run = para.add_run(contract_info)
                        set_font(contract_run, size=Pt(10.5))
                    else:
                        # Just company name, no contract info
                        company_run = para.add_run(company_part)
                        set_font(company_run, size=Pt(10.5), bold=True, color=BLACK)

                    # Pipe separator (now " — ") - 10pt
                    dash_run = para.add_run(' — ')
                    set_font(dash_run, size=Pt(10))

                    # Location - 10pt italic
                    location_part = parts[1].strip() if len(parts) > 1 else ''
                    if location_part:
                        location_run = para.add_run(location_part)
                        set_font(location_run, size=Pt(10), italic=True)

                    # Dates - 9.5pt gray
                    dates_part = parts[2].strip() if len(parts) > 2 else ''
                    if dates_part:
                        dates_run = para.add_run(dates_part)
                        set_font(dates_run, size=Pt(9.5), color=GRAY)

                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)

                else:
                    para = doc.add_paragraph(job_line)
                    run = para.runs[0] if para.runs else para.add_run()
                    set_font(run, size=Pt(10.5), bold=True, color=BLACK)
                    para.paragraph_format.space_after = Pt(3)

            # Bold role title under job - 10pt bold italic #1F3864
            elif item.startswith('**') and item.endswith('**'):
                para = doc.add_paragraph(item.replace('**', ''))
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(10), bold=True, italic=True, color=BLUE)
                para.paragraph_format.space_after = Pt(3)

            # Bullet points (with optional **bold** prefix) - 10pt, List Bullet style
            elif item.startswith('- '):
                bullet_text = item[2:]
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(3.5)

                # Check for **bold** at start of bullet
                if bullet_text.startswith('**'):
                    end_bold = bullet_text.find('**', 2)
                    if end_bold > 0:
                        bold_text = bullet_text[2:end_bold]
                        rest_text = bullet_text[end_bold+2:].strip()
                        bold_run = para.add_run(bold_text + ' ')
                        set_font(bold_run, size=Pt(10), bold=True, color=BLACK)
                        if rest_text:
                            rest_run = para.add_run(rest_text)
                            set_font(rest_run, size=Pt(10), color=BLACK)
                    else:
                        run = para.add_run(bullet_text)
                        set_font(run, size=Pt(10), color=BLACK)
                else:
                    run = para.add_run(bullet_text)
                    set_font(run, size=Pt(10), color=BLACK)

            # Regular bold lines (Education, Clearance, etc.) - 10pt bold black
            elif item.startswith('**'):
                para = doc.add_paragraph(item.replace('**', ''))
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(10), bold=True, color=BLACK)
                para.paragraph_format.space_after = Pt(3)

            # Italic lines - 10pt italic gray
            elif item.startswith('*') and item.endswith('*'):
                para = doc.add_paragraph(item.replace('*', ''))
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(10), italic=True, color=GRAY)
                para.paragraph_format.space_after = Pt(3)

            # URLs - 10pt gray
            elif item.startswith('http'):
                para = doc.add_paragraph(item)
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(10), color=GRAY)
                para.paragraph_format.space_after = Pt(3)

            # Technical expertise lines (Key: Value)
            elif ':' in item and item.split(':')[0].count('**') == 2:
                parts = item.split(':', 1)
                if len(parts) == 2:
                    para = doc.add_paragraph()
                    key_run = para.add_run(parts[0].replace('**', '') + ':')
                    set_font(key_run, size=Pt(10), bold=True, color=BLACK)
                    val_run = para.add_run(parts[1].strip())
                    set_font(val_run, size=Pt(10), color=BLACK)
                    para.paragraph_format.space_after = Pt(3)

            # Default text - 10pt black
            else:
                para = doc.add_paragraph(item)
                run = para.runs[0] if para.runs else para.add_run()
                set_font(run, size=Pt(10), color=BLACK)
                para.paragraph_format.space_after = Pt(3)

def format_document(doc):
    """Apply document-wide formatting."""
    # Set margins: 0.625" L/R, 0.43" T/B, US Letter
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.43)
        section.bottom_margin = Inches(0.43)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)
        section.page_height = Inches(11.0)
        section.page_width = Inches(8.5)

def main():
    if len(sys.argv) != 3:
        print(f"Usage: resume-formatter <input.md> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    # Read markdown
    with open(input_file, 'r') as f:
        md_content = f.read()

    # Parse and build
    sections = parse_markdown(md_content)
    doc = Document()
    format_document(doc)
    build_resume(sections, doc)

    # Save
    doc.save(output_file)
    print(f"Generated: {output_file}")

if __name__ == '__main__':
    main()
